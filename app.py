"""
Doc Studio — App Streamlit para estudo de geração de documentos Word
Bibliotecas: streamlit, docxtpl, python-docx, matplotlib, pandas
"""
import streamlit as st
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import io
import json
import os
import sys
from datetime import date

# Adiciona a pasta raiz ao path para importar utils
sys.path.insert(0, os.path.dirname(__file__))

from utils.extractor import extrair_de_csv, extrair_de_json
from utils.writer import renderizar

st.set_page_config(
    page_title="Doc Studio",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
    .main-header {
        font-size: 2rem;
        font-weight: 700;
        color: #01696f;
        margin-bottom: 0.2rem;
    }
    .sub-header {
        color: #7a7974;
        font-size: 0.95rem;
        margin-bottom: 1.5rem;
    }
    .info-box {
        background: #f0f7f7;
        border-left: 4px solid #01696f;
        padding: 0.8rem 1rem;
        border-radius: 4px;
        margin-bottom: 1rem;
        font-size: 0.9rem;
    }
    .code-label {
        font-family: monospace;
        background: #edeae5;
        padding: 2px 6px;
        border-radius: 3px;
        font-size: 0.85rem;
        color: #28251d;
    }
</style>
""", unsafe_allow_html=True)

st.sidebar.markdown("##  Doc Studio")
st.sidebar.markdown("Projeto de estudo — FGV IBRE")
st.sidebar.divider()

pagina = st.sidebar.radio(
    "Escolha o módulo:",
    [
        "Início",
        "Relatório (formulário)",
        "Relatório (upload CSV)",
        "Contrato",
        "Inspecionar .docx",
        "Referência Jinja2",
    ],
)

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")



if pagina == "Início":
    st.markdown('<div class="main-header">Doc Studio</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Ambiente de estudo para geração e edição de documentos Word com Python</div>', unsafe_allow_html=True)

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### O que você pratica aqui")
        st.markdown("""
        | Módulo | Biblioteca | O que faz |
        |--------|-----------|-----------|
        | Relatório (formulário) | `docxtpl` + `Jinja2` | Preenche template com variáveis diretas |
        | Relatório (CSV) | `pandas` + `docxtpl` | Loop `{%tr for %}` com dados reais |
        | Contrato | `docxtpl` | Condicionais `{% if %}` |
        | Inspecionar .docx | `python-docx` | Lê estrutura interna do Word |
        """)

    with col2:
        st.markdown("### Fluxo do sistema")
        st.markdown("""
        ```
        Formulário / CSV
               ↓
          extractor.py    ← formata dados → dict Python
               ↓
           writer.py      ← renderiza template com docxtpl
               ↓
          template.docx   ← Jinja2 substitui {{ variavel }}
               ↓
         Relatório.docx   ← download pelo usuário
        ```
        """)

    st.info(" **Dica de estudo:** Comece pelo módulo *Relatório (formulário)*, depois tente o CSV, e por fim inspecione o .docx gerado com o módulo *Inspecionar*.")



elif pagina == "Relatório (formulário)":
    st.markdown('<div class="main-header"> Relatório por Formulário</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Preencha os campos — o app gera um .docx com suas variáveis</div>', unsafe_allow_html=True)

    st.markdown('<div class="info-box"> <strong>Conceito:</strong> Cada campo do formulário vira uma <span class="code-label">{{ variavel }}</span> no template Word. O <code>writer.py</code> passa um dicionário Python para o <code>DocxTemplate.render()</code>.</div>', unsafe_allow_html=True)

    with st.form("form_relatorio"):
        col1, col2 = st.columns(2)
        with col1:
            titulo = st.text_input("Título do relatório", "Análise de Desempenho")
            mes_ref = st.text_input("Mês de referência", "Maio/2026")
            responsavel = st.text_input("Responsável", "Ana Silva")
        with col2:
            departamento = st.text_input("Departamento", "Pesquisa e Dados")
            data_geracao = st.date_input("Data de geração", value=date.today()).strftime("%d/%m/%Y")
            observacao = st.text_area("Observação (opcional)", "")

        st.markdown("#### Itens do relatório")
        st.caption("Adicione até 5 itens manualmente")
        itens_raw = []
        for i in range(1, 4):
            c1, c2, c3 = st.columns(3)
            nome = c1.text_input(f"Nome {i}", f"Item {i}", key=f"n{i}")
            cat = c2.text_input(f"Categoria {i}", "Geral", key=f"c{i}")
            val = c3.number_input(f"Valor {i} (R$)", value=1000.0 * i, key=f"v{i}")
            itens_raw.append({"nome": nome, "categoria": cat, "valor": val})

        gerar = st.form_submit_button(" Gerar Documento Word", type="primary")

    if gerar:
        # Formata itens
        itens = []
        for it in itens_raw:
            if it["nome"].strip():
                v = it["valor"]
                itens.append({
                    **it,
                    "valor_fmt": f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", "."),
                })
        total = sum(i["valor"] for i in itens)
        total_fmt = f"R$ {total:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        # Gera gráfico matplotlib
        fig, ax = plt.subplots(figsize=(8, 4))
        nomes = [i["nome"] for i in itens]
        valores = [i["valor"] for i in itens]
        bars = ax.bar(nomes, valores, color="#01696f", edgecolor="white", linewidth=1.5)
        ax.set_title("Distribuição por Item", fontsize=13, fontweight="bold", color="#28251d")
        ax.set_ylabel("Valor (R$)")
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        for bar, v in zip(bars, valores):
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 20,
                    f"R$ {v:,.0f}", ha="center", va="bottom", fontsize=9, color="#28251d")
        plt.tight_layout()

        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
        buf.seek(0)
        grafico_bytes = buf.read()
        plt.close(fig)

        # Exibe preview do gráfico
        st.image(grafico_bytes, caption="Gráfico que será inserido no Word", width=500)

        # Contexto para o template
        contexto = {
            "titulo": titulo,
            "mes_ref": mes_ref,
            "responsavel": responsavel,
            "departamento": departamento,
            "data_geracao": data_geracao,
            "par_intro": f"Este relatório apresenta os resultados de {mes_ref}, "
                         f"consolidados pelo departamento de {departamento}. "
                         f"Foram analisados {len(itens)} itens com valor total de {total_fmt}.",
            "itens": itens,
            "total": total,
            "total_fmt": total_fmt,
            "qtd_itens": len(itens),
            "observacao": observacao,
        }

        try:
            template_path = os.path.join(TEMPLATE_DIR, "relatorio.docx")
            doc_bytes = renderizar(template_path, contexto, graficos={"grafico_barras": grafico_bytes})
            st.success(" Documento gerado com sucesso!")

            st.download_button(
                label=" Baixar Relatório.docx",
                data=doc_bytes,
                file_name=f"relatorio_{mes_ref.replace('/', '-')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:
            st.error(f"Erro ao gerar documento: {e}")
            st.exception(e)

        with st.expander(" Ver dicionário de contexto enviado ao template"):
            ctx_display = {k: v for k, v in contexto.items() if k not in ("grafico_barras",)}
            st.json(ctx_display)



elif pagina == "Relatório (upload CSV)":
    st.markdown('<div class="main-header"> Relatório por Upload CSV</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Faça upload de um CSV e gere o Word com loop Jinja2</div>', unsafe_allow_html=True)

    st.markdown('<div class="info-box"> <strong>Conceito:</strong> O CSV é lido pelo <code>extractor.py</code>, que retorna uma lista de dicts. No template, usamos <span class="code-label">{%tr for item in itens %}</span> para criar uma linha por registro.</div>', unsafe_allow_html=True)

    # CSV de exemplo para download
    csv_exemplo = """nome,valor,categoria
Consultoria Estratégica,15000.00,Serviços
Desenvolvimento de Sistema,28000.00,Tecnologia
Treinamento da Equipe,5500.00,Capacitação
Licença de Software,3200.00,Tecnologia
Relatórios Analíticos,8700.00,Serviços
"""
    st.download_button(
        "Baixar CSV de exemplo",
        data=csv_exemplo,
        file_name="exemplo_dados.csv",
        mime="text/csv",
    )

    uploaded = st.file_uploader("Ou faça upload do seu CSV", type=["csv"])

    col1, col2 = st.columns(2)
    with col1:
        titulo = st.text_input("Título", "Relatório de Custos")
        mes_ref = st.text_input("Mês de referência", "Maio/2026")
    with col2:
        responsavel = st.text_input("Responsável", "Equipe de Dados")
        departamento = st.text_input("Departamento", "Financeiro")

    if uploaded:
        dados = extrair_de_csv(uploaded.read())
        st.success(f"CSV carregado: {dados['qtd_itens']} itens | Total: {dados['total_fmt']}")

        df_preview = pd.DataFrame(dados["itens"])
        st.dataframe(df_preview[["nome", "categoria", "valor_fmt"]], use_container_width=True)

        if st.button("Gerar Relatório Word", type="primary"):
            # Gráfico de pizza por categoria
            cats = df_preview.groupby("categoria")["valor"].sum()
            fig, ax = plt.subplots(figsize=(6, 6))
            ax.pie(cats.values, labels=cats.index, autopct="%1.1f%%",
                   colors=["#01696f", "#4f98a3", "#cedcd8", "#0c4e54", "#317a82"])
            ax.set_title("Distribuição por Categoria", fontweight="bold")
            buf = io.BytesIO()
            fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
            buf.seek(0)
            grafico_bytes = buf.read()
            plt.close(fig)
            st.image(grafico_bytes, width=350)

            total = dados["total"]
            contexto = {
                **dados,
                "titulo": titulo,
                "mes_ref": mes_ref,
                "responsavel": responsavel,
                "departamento": departamento,
                "data_geracao": date.today().strftime("%d/%m/%Y"),
                "par_intro": f"Relatório gerado a partir de dados do CSV. "
                             f"Total de {dados['qtd_itens']} itens registrados, "
                             f"com valor total de {dados['total_fmt']}.",
                "observacao": "",
            }

            try:
                template_path = os.path.join(TEMPLATE_DIR, "relatorio.docx")
                doc_bytes = renderizar(template_path, contexto, graficos={"grafico_barras": grafico_bytes})
                st.download_button(
                    " Baixar Relatório.docx",
                    data=doc_bytes,
                    file_name=f"relatorio_csv_{mes_ref.replace('/', '-')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            except Exception as e:
                st.error(f"Erro: {e}")
                st.exception(e)



elif pagina == "Contrato":
    st.markdown('<div class="main-header"> Gerador de Contrato</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Demonstra condicionais Jinja2 e múltiplas variáveis de texto</div>', unsafe_allow_html=True)

    st.markdown('<div class="info-box"> <strong>Conceito:</strong> O campo <em>Cláusula Extra</em> usa <span class="code-label">{% if clausula_extra %}...{% endif %}</span> — só aparece no Word se preenchido.</div>', unsafe_allow_html=True)

    with st.form("form_contrato"):
        st.subheader("Contratante")
        c1, c2 = st.columns(2)
        contratante_nome = c1.text_input("Nome / Razão Social", "FGV IBRE")
        contratante_doc = c2.text_input("CNPJ", "33.200.049/0001-72")
        contratante_endereco = st.text_input("Endereço", "Rua Barão de Itambi, 60 — Rio de Janeiro/RJ")

        st.subheader("Contratado")
        c3, c4 = st.columns(2)
        contratado_nome = c3.text_input("Nome / Razão Social", "Tech Solutions Ltda.")
        contratado_doc = c4.text_input("CNPJ", "12.345.678/0001-99")

        st.subheader("Objeto e Valores")
        objeto = st.text_area("Descrição do objeto", "Desenvolvimento de sistema de automação de relatórios.")
        c5, c6 = st.columns(2)
        valor_total = c5.number_input("Valor total (R$)", value=50000.0)
        parcelas = c6.number_input("Número de parcelas", value=5, min_value=1)

        st.subheader("Serviços Incluídos")
        servicos = []
        for i in range(1, 4):
            cs1, cs2 = st.columns(2)
            sn = cs1.text_input(f"Serviço {i}", f"Módulo {i}", key=f"sn{i}")
            sp = cs2.text_input(f"Prazo {i}", f"30 dias", key=f"sp{i}")
            if sn.strip():
                servicos.append({"nome": sn, "prazo": sp})

        st.subheader("Datas e Extras")
        c7, c8, c9 = st.columns(3)
        data_inicio = c7.date_input("Início", value=date.today()).strftime("%d/%m/%Y")
        data_fim = c8.date_input("Término").strftime("%d/%m/%Y")
        cidade = c9.text_input("Cidade", "Rio de Janeiro")
        clausula_extra = st.text_area("Cláusula extra (deixe vazio para omitir)", "")

        gerar = st.form_submit_button(" Gerar Contrato Word", type="primary")

    if gerar:
        valor_parcela = valor_total / max(parcelas, 1)
        fmt = lambda v: f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        contexto = {
            "contratante_nome": contratante_nome,
            "contratante_doc": contratante_doc,
            "contratante_endereco": contratante_endereco,
            "contratado_nome": contratado_nome,
            "contratado_doc": contratado_doc,
            "objeto": objeto,
            "valor_total_fmt": fmt(valor_total),
            "parcelas": int(parcelas),
            "valor_parcela_fmt": fmt(valor_parcela),
            "servicos": servicos,
            "data_inicio": data_inicio,
            "data_fim": data_fim,
            "data_assinatura": date.today().strftime("%d de %B de %Y"),
            "cidade": cidade,
            "clausula_extra": clausula_extra,
        }

        try:
            template_path = os.path.join(TEMPLATE_DIR, "contrato.docx")
            doc_bytes = renderizar(template_path, contexto)
            st.success(" Contrato gerado!")
            st.download_button(
                " Baixar Contrato.docx",
                data=doc_bytes,
                file_name="contrato_servicos.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:
            st.error(f"Erro: {e}")
            st.exception(e)

        with st.expander("Ver contexto JSON enviado ao template"):
            st.json(contexto)



elif pagina == "Inspecionar .docx":
    st.markdown('<div class="main-header"> Inspecionar Documento Word</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Use python-docx para ler a estrutura interna de qualquer .docx</div>', unsafe_allow_html=True)

    st.markdown('<div class="info-box"> <strong>Conceito:</strong> Um arquivo <code>.docx</code> é um ZIP contendo XMLs. O <code>python-docx</code> expõe parágrafos, tabelas, runs e estilos como objetos Python — essa é a base do <code>docxtpl</code>.</div>', unsafe_allow_html=True)

    uploaded_docx = st.file_uploader("Faça upload de um .docx para inspecionar", type=["docx"])

    if uploaded_docx:
        from docx import Document as DocxDoc
        doc = DocxDoc(io.BytesIO(uploaded_docx.read()))

        tab1, tab2, tab3 = st.tabs(["Parágrafos", "Tabelas", "Estilos"])

        with tab1:
            st.subheader(f"Parágrafos ({len(doc.paragraphs)})")
            data = []
            for i, p in enumerate(doc.paragraphs):
                data.append({
                    "Índice": i,
                    "Estilo": p.style.name,
                    "Texto (50 chars)": p.text[:80] if p.text else "(vazio)",
                    "Alinhamento": str(p.alignment),
                    "Runs": len(p.runs),
                })
            st.dataframe(pd.DataFrame(data), use_container_width=True)

        with tab2:
            st.subheader(f"Tabelas ({len(doc.tables)})")
            for ti, table in enumerate(doc.tables):
                st.markdown(f"**Tabela {ti + 1}** — {len(table.rows)} linhas × {len(table.columns)} colunas")
                rows_data = []
                for row in table.rows:
                    rows_data.append([cell.text[:40] for cell in row.cells])
                if rows_data:
                    df_t = pd.DataFrame(rows_data)
                    st.dataframe(df_t, use_container_width=True)

        with tab3:
            st.subheader("Estilos usados no documento")
            estilos = list(set(p.style.name for p in doc.paragraphs))
            for e in sorted(estilos):
                st.markdown(f"- `{e}`")

            st.subheader("Imagens embutidas")
            imgs = doc.inline_shapes
            if imgs:
                st.write(f"{len(imgs)} imagem(s) encontrada(s)")
                for i, img in enumerate(imgs):
                    st.write(f"  Imagem {i+1}: {img.width.cm:.1f}cm × {img.height.cm:.1f}cm | Tipo: {img.shape_type}")
            else:
                st.info("Nenhuma imagem embutida encontrada.")
    else:
        st.info("Faça upload de um .docx para começar. Dica: use um dos arquivos gerados nas outras abas!")



elif pagina == "Referência Jinja2":
    st.markdown('<div class="main-header"> Referência Jinja2</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Sintaxe que você precisa dominar para os templates Word</div>', unsafe_allow_html=True)

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### Variáveis")
        st.code("{{ nome }}\n{{ item.valor }}\n{{ item.valor_fmt }}", language="jinja2")

        st.markdown("### Condicionais")
        st.code(
            "{% if observacao %}\n  {{ observacao }}\n{% else %}\n  Sem observações.\n{% endif %}",
            language="jinja2"
        )

        st.markdown("### Filtros úteis")
        st.code(
            "{{ nome | upper }}\n{{ nome | lower }}\n{{ valor | round(2) }}\n{{ lista | length }}",
            language="jinja2"
        )

    with col2:
        st.markdown("### Loop em tabela (docxtpl)")
        st.code(
            "{%tr for item in itens %}\n  {{ item.nome }} | {{ item.valor_fmt }}\n{%tr endfor %}",
            language="jinja2"
        )
        st.caption("Use `{%tr %}` (não `{% %}`) para loops dentro de tabelas Word!")

        st.markdown("### Loop em lista (parágrafo)")
        st.code(
            "{%p for item in itens %}\n  - {{ item.nome }}: {{ item.valor_fmt }}\n{%p endfor %}",
            language="jinja2"
        )

        st.markdown("### Acessar índice")
        st.code("{% for item in itens %}\n  {{ loop.index }}. {{ item.nome }}\n{% endfor %}", language="jinja2")

    st.divider()
    st.markdown("### Teste de sintaxe Jinja2 ao vivo")
    st.caption("Cole um template e um contexto JSON para testar antes de usar no Word")

    col3, col4 = st.columns(2)
    with col3:
        template_txt = st.text_area(
            "Template Jinja2",
            "Olá, {{ nome }}!\nTotal: {{ total_fmt }}\n{% if nota %}Nota: {{ nota }}{% endif %}",
            height=150,
        )
    with col4:
        contexto_txt = st.text_area(
            "Contexto JSON",
            '{"nome": "Maria", "total_fmt": "R$ 1.500,00", "nota": "Aprovado"}',
            height=150,
        )

    if st.button("Renderizar"):
        try:
            from jinja2 import Environment
            env = Environment()
            tmpl = env.from_string(template_txt)
            ctx = json.loads(contexto_txt)
            resultado = tmpl.render(ctx)
            st.success("Resultado:")
            st.code(resultado)
        except Exception as e:
            st.error(f"Erro: {e}")

